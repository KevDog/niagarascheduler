"""
Word Template Processing Service

Handles Word document template processing with variable replacement
using the NU 2025 Syllabus Template.docx
"""

import os
import re
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches
from tempfile import NamedTemporaryFile


class WordTemplateService:
    """Service for processing Word document templates"""
    
    def __init__(self, template_dir: str):
        self.template_dir = template_dir
        self.template_path = os.path.join(template_dir, 'NU 2025 Syllabus Template.docx')
    
    def process_template(self, template_data: Dict[str, Any]) -> str:
        """
        Process Word template by replacing placeholders with actual data
        
        Args:
            template_data: Dictionary containing template variables and their values
            
        Returns:
            Path to the processed Word document
            
        Raises:
            FileNotFoundError: If template file doesn't exist
            Exception: If processing fails
        """
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f'Template file not found: {self.template_path}')
        
        try:
            # Load the Word document template
            doc = Document(self.template_path)
            
            # Replace placeholders in paragraphs
            self._replace_placeholders_in_paragraphs(doc, template_data)
            
            # Replace placeholders in tables
            self._replace_placeholders_in_tables(doc, template_data)
            
            # Replace placeholders in headers and footers
            self._replace_placeholders_in_headers_footers(doc, template_data)
            
            # Create temporary file for the processed document
            temp_file = NamedTemporaryFile(suffix='.docx', delete=False)
            doc.save(temp_file.name)
            temp_file.close()
            
            return temp_file.name
            
        except Exception as e:
            raise Exception(f'Error processing Word template: {str(e)}')
    
    def _replace_placeholders_in_paragraphs(self, doc: Document, template_data: Dict[str, Any]):
        """Replace placeholders in document paragraphs"""
        for paragraph in doc.paragraphs:
            self._replace_text_in_paragraph(paragraph, template_data)
    
    def _replace_placeholders_in_tables(self, doc: Document, template_data: Dict[str, Any]):
        """Replace placeholders in document tables"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_text_in_paragraph(paragraph, template_data)
    
    def _replace_placeholders_in_headers_footers(self, doc: Document, template_data: Dict[str, Any]):
        """Replace placeholders in headers and footers"""
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    self._replace_text_in_paragraph(paragraph, template_data)
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    self._replace_text_in_paragraph(paragraph, template_data)
    
    def _replace_text_in_paragraph(self, paragraph, template_data: Dict[str, Any]):
        """Replace template variables in a paragraph while preserving formatting"""
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Find all placeholders in the format {{VARIABLE_NAME}}
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)
        
        if not placeholders:
            return
        
        # Replace each placeholder
        new_text = full_text
        for placeholder in placeholders:
            # Clean up the placeholder name (remove extra spaces)
            clean_placeholder = placeholder.strip()
            
            # Get the replacement value
            replacement_value = template_data.get(clean_placeholder, f'{{{{{clean_placeholder}}}}}')
            
            # Convert to string and handle None values
            if replacement_value is None:
                replacement_value = ''
            elif not isinstance(replacement_value, str):
                replacement_value = str(replacement_value)
            
            # Replace the placeholder with the actual value
            new_text = new_text.replace(f'{{{{{placeholder}}}}}', replacement_value)
        
        # If text changed, update the paragraph
        if new_text != full_text:
            # Clear existing runs and set new text
            paragraph.clear()
            paragraph.add_run(new_text)
    
    def prepare_template_data(self, schedule_data: List[str], semester: str, year: str,
                            course_id: str = '', instructor_name: str = 'TBD',
                            office_hours: str = '', office_location: str = '',
                            email_address: str = '', phone_number: str = '',
                            textbooks: str = '', assignments: str = '',
                            attendance_policy: str = '', grading_policy: str = '',
                            ai_policy: str = '', bibliography: str = '',
                            course_description: str = '', course_title: str = '',
                            **kwargs) -> Dict[str, Any]:
        """
        Prepare template data dictionary for Word document replacement
        
        Returns:
            Dictionary with template variables mapped to their values
        """
        
        # Format semester display
        semester_parts = semester.split('_')
        semester_year = f"20{semester_parts[0]}" if len(semester_parts) > 0 else str(year)
        semester_name = semester_parts[1] if len(semester_parts) > 1 else 'Unknown'
        semester_display = f"{semester_name} {semester_year}"
        
        # Prepare schedule table content
        schedule_table = self._format_schedule_for_template(schedule_data)
        
        # Department mission statement (placeholder - could be enhanced to load from data)
        department_code = course_id.split()[0] if ' ' in course_id else ''
        department_mission = kwargs.get('department_mission_statement', 
            'This department is committed to providing excellent education and preparing students for success.')
        
        return {
            'COURSE_ID': course_id or 'Course ID',
            'COURSE_TITLE': course_title or f'Course {course_id}' if course_id else 'Course Title',
            'SEMESTER': semester_display,
            'YEAR': str(year),
            'INSTRUCTOR_NAME': instructor_name or 'TBD',
            'OFFICE_HOURS': office_hours or 'Office hours will be announced.',
            'OFFICE_LOCATION': office_location or 'Office location will be announced.',
            'EMAIL_ADDRESS': email_address or 'Email will be provided.',
            'PHONE_NUMBER': phone_number or 'Phone number will be provided.',
            'COURSE_DESCRIPTION': course_description or 'Course description will be provided.',
            'DEPARTMENT_MISSION_STATEMENT': department_mission,
            'SCHEDULE_TABLE': schedule_table,
            'TEXTBOOKS': textbooks or 'Textbooks will be announced.',
            'ASSIGNMENTS': assignments or 'Assignment details will be provided.',
            'ATTENDANCE_POLICY': attendance_policy or 'Regular attendance is expected.',
            'GRADING_POLICY': grading_policy or 'Grading criteria will be provided.',
            'AI_POLICY': ai_policy or 'AI policy will be specified.',
            'BIBLIOGRAPHY': bibliography or 'References will be provided as needed.'
        }
    
    def _format_schedule_for_template(self, schedule_data: List[str]) -> str:
        """Format schedule data for insertion into Word template"""
        if not schedule_data:
            return 'Schedule will be provided.'
        
        # Convert schedule data to a formatted string
        # This assumes schedule_data contains formatted date/event information
        schedule_lines = []
        for item in schedule_data:
            if isinstance(item, dict):
                date = item.get('date', '')
                event_type = item.get('type', '')
                name = item.get('name', '')
                
                if event_type == 'class':
                    schedule_lines.append(f"{date} - Class Meeting")
                elif event_type == 'holiday':
                    schedule_lines.append(f"{date} - {name}")
                elif event_type == 'event':
                    schedule_lines.append(f"{date} - {name}")
                else:
                    schedule_lines.append(f"{date} - {name}")
            else:
                # Handle string entries
                schedule_lines.append(str(item))
        
        return '\n'.join(schedule_lines) if schedule_lines else 'Schedule will be provided.'


def create_word_syllabus(schedule_data: List[str], semester: str, year: str,
                        template_dir: str, **kwargs) -> str:
    """
    Convenience function to create a Word syllabus from template
    
    Args:
        schedule_data: Schedule data
        semester: Semester
        year: Year
        template_dir: Directory containing the template
        **kwargs: Additional template data
    
    Returns:
        Path to the generated Word document
    """
    service = WordTemplateService(template_dir)
    template_data = service.prepare_template_data(
        schedule_data=schedule_data,
        semester=semester,
        year=year,
        **kwargs
    )
    return service.process_template(template_data)